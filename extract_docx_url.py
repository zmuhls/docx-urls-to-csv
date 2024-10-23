import csv
import re
import osopen 
from docx import Document
import requests

# Function to extract URLs from the text
def extract_urls(text):
    url_pattern = re.compile(r'(https?://[^\s.]+(?:\.[^\s.]+)*)')
    return url_pattern.findall(text)

# Function to extract text from a Word document
def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

# Function to extract the first heading (h1) from a Word document
def extract_first_heading(docx_file):
    doc = Document(docx_file)
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading 1'):
            return paragraph.text
    return "No Heading 1 Found"

# Function to check the status of a URL
def check_url_status(url):
    try:
        response = requests.head(url, allow_redirects=True, timeout=5)
        return response.status_code
    except requests.RequestException as e:
        return f"Error: {e}"

# Function to save URLs and their status to a CSV file
def save_urls_to_csv(urls, csv_file):
    with open(csv_file, mode='w', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        writer.writerow(["URL", "Status", "Source Document", "Document Name"])  # Header
        for url, source_doc, doc_name in urls:
            status = check_url_status(url)
            writer.writerow([url, status, source_doc, doc_name])

# Function to find all .docx files in the current working directory
def find_docx_files():
    return [file for file in os.listdir() if file.endswith('.docx')]

# Function to concatenate all .docx files into one document
def concatenate_docx_files(output_docx_file):
    docx_files = find_docx_files()
    combined_doc = Document()

    for docx_file in docx_files:
        print(f"Adding {docx_file} to the combined document...")
        try:
            doc = Document(docx_file)
            for paragraph in doc.paragraphs:
                combined_doc.add_paragraph(paragraph.text)
        except Exception as e:
            print(f"Error processing {docx_file}: {e}")

    combined_doc.save(output_docx_file)
    print(f"All documents concatenated into {output_docx_file}")

# Main function to process all Word documents, concatenate them, and output URLs to a CSV
def extract_urls_from_concatenated_docx_to_csv(output_docx_file, output_csv_file):
    concatenate_docx_files(output_docx_file)
    docx_files = find_docx_files()
    all_urls = []

    for docx_file in docx_files:
        text = extract_text_from_docx(docx_file)
        urls = extract_urls(text)
        doc_name = extract_first_heading(docx_file)
        all_urls.extend([(url, docx_file, doc_name) for url in urls])

    save_urls_to_csv(all_urls, output_csv_file)
    print(f"URLs extracted and saved to {output_csv_file}")

# Usage example
output_docx_file = 'combined_document.docx'  # Output concatenated DOCX file
output_csv_file = 'extracted_urls.csv'  # Output CSV file
extract_urls_from_concatenated_docx_to_csv(output_docx_file, output_csv_file)

print(f"Processing complete. All URLs and their statuses are saved in {output_csv_file}.")